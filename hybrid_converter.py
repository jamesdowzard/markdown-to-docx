import pypandoc
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os
import tempfile
import re

class HybridMarkdownConverter:
    def __init__(self, template_path=None):
        self.template_path = template_path
        
    def convert(self, markdown_file_path, output_docx_path, enhance_formatting=True):
        """
        Convert markdown to docx using pypandoc, then enhance with python-docx
        """
        try:
            # First, convert using pypandoc
            temp_docx = self._convert_with_pypandoc(markdown_file_path)
            
            if enhance_formatting:
                # Then enhance the formatting with python-docx
                self._enhance_formatting(temp_docx, output_docx_path)
            else:
                # Just copy the pypandoc output
                import shutil
                shutil.copy(temp_docx, output_docx_path)
            
            # Clean up temporary file
            os.unlink(temp_docx)
            
            print(f"Successfully converted '{markdown_file_path}' to '{output_docx_path}'")
            
        except Exception as e:
            print(f"Error converting '{markdown_file_path}': {e}")
            raise
    
    def _convert_with_pypandoc(self, markdown_file_path):
        """
        Use pypandoc for initial conversion
        """
        # Create temporary output file
        temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
        os.close(temp_fd)
        
        # Configure pypandoc options for better Word output
        extra_args = [
            '--reference-doc=' + self.template_path if self.template_path else None,
            '--wrap=none',  # Don't wrap lines
            '--standalone',  # Produce standalone document
            '--table-of-contents' if self._has_headings(markdown_file_path) else None,
        ]
        
        # Filter out None values
        extra_args = [arg for arg in extra_args if arg is not None]
        
        # Convert markdown to docx
        pypandoc.convert_file(
            markdown_file_path,
            'docx',
            outputfile=temp_path,
            extra_args=extra_args
        )
        
        return temp_path
    
    def _has_headings(self, markdown_file_path):
        """
        Check if markdown file has headings (for TOC generation)
        """
        with open(markdown_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return bool(re.search(r'^#+\s+', content, re.MULTILINE))
    
    def _enhance_formatting(self, input_docx_path, output_docx_path):
        """
        Enhance the pypandoc output with python-docx
        """
        doc = Document(input_docx_path)
        
        # Enhance tables
        self._enhance_tables(doc)
        
        # Enhance code blocks
        self._enhance_code_blocks(doc)
        
        # Enhance lists
        self._enhance_lists(doc)
        
        # Save the enhanced document
        doc.save(output_docx_path)
    
    def _enhance_tables(self, doc):
        """
        Add borders and improve table formatting
        """
        for table in doc.tables:
            # Set table alignment
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add borders to all cells
            for row in table.rows:
                for cell in row.cells:
                    self._add_cell_border(cell)
                    
                    # Make header row bold
                    if row == table.rows[0]:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def _add_cell_border(self, cell):
        """
        Add border to a table cell
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create borders element
        tcBorders = OxmlElement('w:tcBorders')
        
        # Add all borders
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_element = OxmlElement(f'w:{border_name}')
            border_element.set(qn('w:val'), 'single')
            border_element.set(qn('w:sz'), '4')
            border_element.set(qn('w:space'), '0')
            border_element.set(qn('w:color'), '000000')
            tcBorders.append(border_element)
        
        tcPr.append(tcBorders)
    
    def _enhance_code_blocks(self, doc):
        """
        Improve code block formatting
        """
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Code' or 'code' in paragraph.style.name.lower():
                # Set monospace font
                for run in paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                
                # Add background color (light gray)
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(6)
    
    def _enhance_lists(self, doc):
        """
        Improve list formatting
        """
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('List'):
                # Improve list spacing
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)

class AdvancedMarkdownConverter:
    """
    Extended converter with custom parsing for advanced markdown elements
    """
    
    def __init__(self, template_path=None):
        self.template_path = template_path
        if template_path:
            self.document = Document(template_path)
        else:
            self.document = Document()
    
    def convert(self, markdown_file_path, output_docx_path):
        """
        Convert markdown with full custom parsing
        """
        with open(markdown_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Parse and convert all elements
        self._parse_content(content)
        
        # Save document
        self.document.save(output_docx_path)
        print(f"Successfully converted '{markdown_file_path}' to '{output_docx_path}'")
    
    def _parse_content(self, content):
        """
        Parse markdown content with support for all elements
        """
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Skip empty lines
            if not line.strip():
                i += 1
                continue
            
            # Parse different elements
            if self._is_heading(line):
                self._add_heading(line)
            elif self._is_code_block_start(line):
                i = self._parse_code_block(lines, i)
            elif self._is_table_line(line):
                i = self._parse_table(lines, i)
            elif self._is_blockquote(line):
                i = self._parse_blockquote(lines, i)
            elif self._is_list_item(line):
                i = self._parse_list(lines, i)
            elif self._is_task_list(line):
                i = self._parse_task_list(lines, i)
            elif self._is_horizontal_rule(line):
                self._add_horizontal_rule()
            else:
                self._add_paragraph(line)
            
            i += 1
    
    def _is_heading(self, line):
        return line.strip().startswith('#')
    
    def _is_code_block_start(self, line):
        return line.strip().startswith('```')
    
    def _is_table_line(self, line):
        return '|' in line and line.strip().startswith('|')
    
    def _is_blockquote(self, line):
        return line.strip().startswith('>')
    
    def _is_list_item(self, line):
        return re.match(r'^\s*[\*\-\+]\s+', line) or re.match(r'^\s*\d+\.\s+', line)
    
    def _is_task_list(self, line):
        return re.match(r'^\s*[\*\-\+]\s+\[[\sx]\]\s+', line)
    
    def _is_horizontal_rule(self, line):
        return re.match(r'^\s*[\*\-_]{3,}\s*$', line)
    
    def _add_heading(self, line):
        match = re.match(r'^(#+)\s*(.*)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            self.document.add_heading(text, level=min(level, 6))
    
    def _parse_code_block(self, lines, start_index):
        """Parse code block and return end index"""
        code_lines = []
        i = start_index + 1
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        code_content = '\n'.join(code_lines)
        p = self.document.add_paragraph()
        run = p.add_run(code_content)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        return i
    
    def _parse_table(self, lines, start_index):
        """Parse table and return end index"""
        table_lines = []
        i = start_index
        
        # Collect all table lines
        while i < len(lines) and self._is_table_line(lines[i]):
            table_lines.append(lines[i])
            i += 1
        
        if len(table_lines) >= 2:
            # Parse header
            header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
            
            # Skip separator line
            data_lines = table_lines[2:] if len(table_lines) > 2 else []
            
            # Create table
            table = self.document.add_table(rows=1, cols=len(header_cells))
            
            # Add header
            for j, header in enumerate(header_cells):
                table.cell(0, j).text = header
                # Make header bold
                for paragraph in table.cell(0, j).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add data rows
            for line in data_lines:
                data_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(data_cells) == len(header_cells):
                    row = table.add_row()
                    for j, data in enumerate(data_cells):
                        row.cells[j].text = data
        
        return i - 1
    
    def _parse_blockquote(self, lines, start_index):
        """Parse blockquote and return end index"""
        quote_lines = []
        i = start_index
        
        while i < len(lines) and lines[i].strip().startswith('>'):
            quote_line = lines[i].strip()[1:].strip()
            quote_lines.append(quote_line)
            i += 1
        
        quote_text = ' '.join(quote_lines)
        p = self.document.add_paragraph()
        p.style = 'Quote'
        p.add_run(quote_text)
        
        return i - 1
    
    def _parse_list(self, lines, start_index):
        """Parse list and return end index"""
        i = start_index
        
        while i < len(lines) and self._is_list_item(lines[i]):
            line = lines[i]
            
            # Determine list type and content
            if re.match(r'^\s*\d+\.\s+', line):
                # Ordered list
                content = re.sub(r'^\s*\d+\.\s+', '', line)
                p = self.document.add_paragraph(style='List Number')
            else:
                # Unordered list
                content = re.sub(r'^\s*[\*\-\+]\s+', '', line)
                p = self.document.add_paragraph(style='List Bullet')
            
            self._add_formatted_text(p, content)
            i += 1
        
        return i - 1
    
    def _parse_task_list(self, lines, start_index):
        """Parse task list and return end index"""
        i = start_index
        
        while i < len(lines) and self._is_task_list(lines[i]):
            line = lines[i]
            
            # Check if task is completed
            is_completed = '[x]' in line or '[X]' in line
            
            # Extract task content
            content = re.sub(r'^\s*[\*\-\+]\s+\[[\sx]\]\s+', '', line)
            
            p = self.document.add_paragraph(style='List Bullet')
            checkbox = '☑' if is_completed else '☐'
            p.add_run(f"{checkbox} ")
            self._add_formatted_text(p, content)
            
            i += 1
        
        return i - 1
    
    def _add_horizontal_rule(self):
        """Add horizontal rule"""
        p = self.document.add_paragraph()
        p.add_run('_' * 50)
    
    def _add_paragraph(self, text):
        """Add paragraph with formatting"""
        p = self.document.add_paragraph()
        self._add_formatted_text(p, text)
    
    def _add_formatted_text(self, paragraph, text):
        """Add text with inline formatting (bold, italic, code, etc.)"""
        # Handle links first
        text = self._process_links(paragraph, text)
        
        # Handle inline formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|~~.*?~~)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            elif part.startswith('~~') and part.endswith('~~'):
                run = paragraph.add_run(part[2:-2])
                run.font.strike = True
            else:
                paragraph.add_run(part)
    
    def _process_links(self, paragraph, text):
        """Process markdown links"""
        # Simple link processing - could be enhanced
        link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
        
        def replace_link(match):
            link_text = match.group(1)
            url = match.group(2)
            
            # Add hyperlink (simplified - Word hyperlinks are complex)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = None  # Default blue color
            run.underline = True
            
            return ""  # Return empty string to remove from text
        
        return re.sub(link_pattern, replace_link, text)


if __name__ == "__main__":
    # Test hybrid converter
    hybrid_converter = HybridMarkdownConverter()
    
    # Convert test files
    hybrid_converter.convert(
        "test_basic.md",
        "outputs/test_basic_hybrid.docx"
    )
    
    hybrid_converter.convert(
        "test_complex.md", 
        "outputs/test_complex_hybrid.docx"
    )
    
    print("Hybrid conversion completed!")
    
    # Test advanced converter
    advanced_converter = AdvancedMarkdownConverter()
    
    advanced_converter.convert(
        "test_basic.md",
        "outputs/test_basic_advanced.docx"
    )
    
    advanced_converter.convert(
        "test_complex.md",
        "outputs/test_complex_advanced.docx"
    )
    
    print("Advanced conversion completed!")