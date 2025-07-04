from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class MarkdownToDocxConverter:
    def __init__(self, template_path=None):
        if template_path:
            self.document = Document(template_path)
        else:
            self.document = Document()

    def _add_heading(self, text, level):
        # Map Markdown heading levels to Word's built-in heading styles
        # Assuming styles like 'Heading 1', 'Heading 2', etc., exist in the template or default document
        if 1 <= level <= 9: # Word supports up to Heading 9
            self.document.add_heading(text, level=level)
        else:
            self.document.add_paragraph(text, style='Normal') # Fallback for unsupported levels

    def _add_paragraph(self, text):
        p = self.document.add_paragraph()
        self._add_text_with_formatting(p, text)

    def _add_text_with_formatting(self, paragraph, text):
        # Handle bold, italic, and inline code
        # This is a simplified approach and might not cover all edge cases
        parts = re.split(r'(\**.*?\**|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New' # Or a suitable monospace font
            else:
                paragraph.add_run(part)

    def _add_list_item(self, text, level, ordered=False):
        # Use 'List Paragraph' style and handle indentation manually for now
        # For true Word list styles, a template with defined 'List Bullet'/'List Number' styles is best
        paragraph = self.document.add_paragraph(style='List Paragraph')
        paragraph.paragraph_format.left_indent = Inches(0.25 * level)
        if ordered:
            # This is a basic way to add numbers, not a true Word ordered list style
            paragraph.add_run(f'{level}. ')
        else:
            # This is a basic way to add bullets, not a true Word bullet list style
            paragraph.add_run('• ')
        self._add_text_with_formatting(paragraph, text)

    def _add_code_block(self, text):
        try:
            self.document.add_paragraph(text, style='Code')
        except:
            # Fallback if 'Code' style doesn't exist
            p = self.document.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

    def _add_table(self, header, data):
        table = self.document.add_table(rows=1 + len(data), cols=len(header))
        table.style = 'Table Grid' # Assuming 'Table Grid' style exists in template

        # Add header row
        hdr_cells = table.rows[0].cells
        for i, h_text in enumerate(header):
            hdr_cells[i].text = h_text
            # Optional: make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                row_cells[i].text = cell_text

    def convert(self, markdown_file_path, output_docx_path):
        with open(markdown_file_path, 'r') as f:
            lines = f.readlines()

        in_code_block = False
        code_block_content = []
        table_header = []
        table_data = []
        in_table = False

        for line in lines:
            line = line.strip()

            # Code blocks
            if line.startswith("```"):
                if in_code_block:
                    self._add_code_block("\n".join(code_block_content))
                    code_block_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_block_content.append(line)
                continue

            # Tables
            # Check for table header (first line of a table, not inside a code block)
            if re.match(r'^\|.*\|\s*$', line) and not in_table and not in_code_block:
                # This is likely the header line
                table_header = [h.strip() for h in line.split('|')[1:-1]]
                in_table = True
                continue
            # Check for separator line (second line of a table)
            elif re.match(r'^\|[ :\-]+\|[ :\-]+\|.*\|\s*$', line) and in_table and not in_code_block:
                # This is the separator line, ignore it
                continue
            # Check for data row (subsequent lines of a table)
            elif re.match(r'^\|.*\|\s*$', line) and in_table and not in_code_block:
                # This is a data row
                table_data.append([d.strip() for d in line.split('|')[1:-1]])
                continue
            elif in_table and not line.strip() and not in_code_block: # Empty line after table
                # End of table
                self._add_table(table_header, table_data)
                table_header = []
                table_data = []
                in_table = False
                continue # Process the next line
            elif in_table and not in_code_block and not re.match(r'^\|.*\|\s*$', line): # Line not part of table
                # End of table if the line doesn't match table format
                self._add_table(table_header, table_data)
                table_header = []
                table_data = []
                in_table = False

            # Headings
            heading_match = re.match(r'^(#+)\s*(.*)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                self._add_heading(text, level)
                continue

            # Unordered lists
            if re.match(r'^[\*\-+]\s+(.*)$', line):
                # Simple bullet list, assuming level 1 for now
                self._add_list_item(re.match(r'^[\*\-+]\s+(.*)$', line).group(1).strip(), 1, ordered=False)
                continue

            # Ordered lists
            if re.match(r'^\d+\.\s+(.*)$', line):
                # Simple ordered list, assuming level 1 for now
                self._add_list_item(re.match(r'^\d+\.\s+(.*)$', line).group(1).strip(), 1, ordered=True)
                continue

            # Horizontal Rule (--- or *** or ___)
            if re.match(r'^[\-\*_]{3,}$', line):
                # python-docx doesn't have a direct horizontal rule. Add a paragraph with a border.
                p = self.document.add_paragraph()
                p.add_run().add_break()
                continue

            # Paragraphs
            if line:
                self._add_paragraph(line)

        # Handle any pending table or code block at the end of the file
        if in_table:
            self._add_table(table_header, table_data)
        if in_code_block:
            self._add_code_block("\n".join(code_block_content))

        self.document.save(output_docx_path)
        print(f"Successfully converted '{markdown_file_path}' to '{output_docx_path}'")

if __name__ == "__main__":
    # IMPORTANT: Create a 'template.docx' file in the same directory as this script.
    # In Word, create a new blank document. Define/ensure the following styles exist:
    # - Heading 1, Heading 2, ..., Heading 6
    # - List Paragraph (for lists)
    # - Code (for code blocks - you might need to create this style manually, e.g., Courier New font)
    # - Table Grid (for tables)
    # Save this blank document as 'template.docx'.
    
    converter = MarkdownToDocxConverter(template_path="template.docx")
    
    # Convert the basic test file
    converter.convert(
        "test_basic.md",
        "outputs/test_basic_styled.docx"
    )

    # Convert the complex test file
    converter.convert(
        "test_complex.md",
        "outputs/test_complex_styled.docx"
    )